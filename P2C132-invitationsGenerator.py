#! python
import copy
import docx, os

FOLDER_NAME = "P2C13"

os.makedirs(FOLDER_NAME, exist_ok=True)
os.chdir(FOLDER_NAME)

doc = docx.Document("invitation_template.docx")

print("Opening the guest list..")
guestFile = open("guest_list.txt")
guestList = guestFile.readlines()

inviteTemplate = copy.copy(doc.paragraphs)

print("Adding guest names to invitation...")
for guest in guestList:
    doc.add_page_break()
    for i in range(len(inviteTemplate)):
        if inviteTemplate[i].text == "Invitee":
            doc.add_paragraph(guest.strip(), inviteTemplate[i].style)
        else:
            doc.add_paragraph(inviteTemplate[i].text, inviteTemplate[i].style)

print("Saving invitation....")
doc.save("updated_invitation.docx")

print("Done")
